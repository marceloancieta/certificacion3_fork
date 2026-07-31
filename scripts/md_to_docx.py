# -*- coding: utf-8 -*-
"""Convert retos/RETOS_COMPLETOS.md to a Word .docx."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
MD = ROOT / "retos" / "RETOS_COMPLETOS.md"
OUT = ROOT / "retos" / "RETOS_COMPLETOS.docx"


def set_run_font(run, *, size=11, bold=False, italic=False, code=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = "Consolas" if code else "Calibri"
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), "Consolas" if code else "Calibri")
    rFonts.set(qn("w:hAnsi"), "Consolas" if code else "Calibri")
    if color is not None:
        run.font.color.rgb = color


def add_inline_runs(paragraph, text: str, *, size=11):
    """Parse **bold**, `code`, and plain text into runs."""
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos : m.start()])
            set_run_font(run, size=size)
        token = m.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, bold=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=size, code=True, color=RGBColor(0x1A, 0x1A, 0x1A))
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=size)


def add_paragraph(doc, text: str, *, style=None, size=11, space_after=6):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    add_inline_runs(p, text, size=size)
    return p


def parse_table_row(line: str) -> list[str]:
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return cells


def is_table_sep(line: str) -> bool:
    return bool(re.match(r"^\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$", line))


def convert(md_text: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    lines = md_text.splitlines()
    i = 0
    in_code = False
    code_buf: list[str] = []

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_buf = []
            else:
                in_code = False
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(8)
                run = p.add_run("\n".join(code_buf))
                set_run_font(run, size=9, code=True)
                p.paragraph_format.left_indent = Inches(0.15)
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        if line.startswith("# "):
            p = doc.add_heading(line[2:].strip(), level=1)
            for run in p.runs:
                set_run_font(run, size=18, bold=True, color=RGBColor(0x1B, 0x2A, 0x4A))
            i += 1
            continue

        if line.startswith("## "):
            p = doc.add_heading(line[3:].strip(), level=2)
            for run in p.runs:
                set_run_font(run, size=14, bold=True, color=RGBColor(0x0E, 0x7C, 0x7B))
            i += 1
            continue

        if line.startswith("### "):
            p = doc.add_heading(line[4:].strip(), level=3)
            for run in p.runs:
                set_run_font(run, size=12, bold=True)
            i += 1
            continue

        if line.strip() == "---":
            p = doc.add_paragraph()
            run = p.add_run("─" * 48)
            set_run_font(run, size=10, color=RGBColor(0x99, 0x99, 0x99))
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(10)
            i += 1
            continue

        if line.startswith("> "):
            text = line[2:].strip()
            p = add_paragraph(doc, text, size=10, space_after=2)
            p.paragraph_format.left_indent = Inches(0.2)
            for run in p.runs:
                run.italic = True
            i += 1
            continue

        if line.startswith("| ") or (line.startswith("|") and "|" in line[1:]):
            rows: list[list[str]] = []
            while i < len(lines) and "|" in lines[i]:
                if is_table_sep(lines[i]):
                    i += 1
                    continue
                rows.append(parse_table_row(lines[i]))
                i += 1
            if rows:
                cols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=cols)
                table.style = "Table Grid"
                for r_idx, row in enumerate(rows):
                    for c_idx in range(cols):
                        cell_text = row[c_idx] if c_idx < len(row) else ""
                        cell = table.rows[r_idx].cells[c_idx]
                        cell.text = ""
                        p = cell.paragraphs[0]
                        add_inline_runs(p, cell_text, size=9 if r_idx else 9)
                        if r_idx == 0:
                            for run in p.runs:
                                run.bold = True
                doc.add_paragraph()
            continue

        if re.match(r"^- \[[ xX]\] ", line):
            text = re.sub(r"^- \[[ xX]\] ", "", line)
            add_paragraph(doc, "☐ " + text, size=11)
            i += 1
            continue

        if line.startswith("- "):
            text = line[2:]
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            add_inline_runs(p, text, size=11)
            i += 1
            continue

        if re.match(r"^\d+\.\s", line):
            text = re.sub(r"^\d+\.\s+", "", line)
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(3)
            add_inline_runs(p, text, size=11)
            i += 1
            continue

        # Question lines like **1.** ...
        add_paragraph(doc, line.strip(), size=11, space_after=4)
        i += 1

    return doc


def main() -> None:
    md = MD.read_text(encoding="utf-8")
    doc = convert(md)
    doc.save(OUT)
    print(f"Generado: {OUT} ({OUT.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
